"""
Módulo de geração do documento DOCX no formato "Movimentação Bancária".
Replica o modelo em branco que a mãe do usuário usa, com créditos em negrito.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def _set_cell_border(cell, **kwargs):
    """Define bordas em uma célula da tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')

    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)

    tcPr.append(tcBorders)


def _formatar_celula(cell, texto: str, negrito: bool = False, tamanho: int = 9,
                      alinhamento=WD_ALIGN_PARAGRAPH.LEFT):
    """Formata o texto de uma célula com as propriedades desejadas."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alinhamento

    # Reduzir espaçamento do parágrafo
    pf = paragraph.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = Pt(11)

    run = paragraph.add_run(texto)
    run.font.size = Pt(tamanho)
    run.font.name = "Arial"
    run.font.bold = negrito

    # Garantir fonte Arial no XML
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/>')
    rPr.insert(0, rFonts)


def _extrair_info_extrato(lancamentos: list[dict]) -> dict:
    """Extrai informações do extrato (ano, mês) a partir dos lançamentos."""
    if not lancamentos:
        return {"ano": "2022", "mes": ""}

    primeira_data = lancamentos[0].get("data", "")
    partes = primeira_data.split("/")
    if len(partes) == 3:
        return {"ano": partes[2], "mes": partes[1]}
    return {"ano": "2022", "mes": ""}


MESES = {
    "01": "JANEIRO", "02": "FEVEREIRO", "03": "MARÇO", "04": "ABRIL",
    "05": "MAIO", "06": "JUNHO", "07": "JULHO", "08": "AGOSTO",
    "09": "SETEMBRO", "10": "OUTUBRO", "11": "NOVEMBRO", "12": "DEZEMBRO"
}


def gerar_docx(lancamentos: list[dict], output_path: str,
               agencia: str = "3050", conta: str = "7223-0",
               banco: str = "BRADESCO") -> str:
    """
    Gera o documento DOCX no formato da movimentação bancária.

    Args:
        lancamentos: Lista de dicts com data, tipo, descricao, valor
        output_path: Caminho para salvar o DOCX
        agencia: Número da agência
        conta: Número da conta
        banco: Nome do banco

    Returns:
        Caminho do arquivo gerado
    """
    doc = Document()

    # Configurar margens da página
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    info = _extrair_info_extrato(lancamentos)
    ano = info["ano"]

    # Agrupar lançamentos por mês
    lancamentos_por_mes = {}
    for lanc in lancamentos:
        data = lanc.get("data", "")
        partes = data.split("/")
        if len(partes) == 3:
            mes = partes[1]
        else:
            mes = "00"
        if mes not in lancamentos_por_mes:
            lancamentos_por_mes[mes] = []
        lancamentos_por_mes[mes].append(lanc)

    # Ordenar meses
    meses_ordenados = sorted(lancamentos_por_mes.keys())

    primeiro_mes = True
    for mes in meses_ordenados:
        lancs_mes = lancamentos_por_mes[mes]
        nome_mes = MESES.get(mes, f"MÊS {mes}")

        if not primeiro_mes:
            doc.add_page_break()
        primeiro_mes = False

        # Título
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = titulo.add_run(f"MOVIMENTAÇÃO BANCÁRIA ANO {ano}")
        run_titulo.font.size = Pt(14)
        run_titulo.font.bold = True
        run_titulo.font.name = "Arial"
        run_titulo.font.color.rgb = RGBColor(0, 0, 0)

        # Subtítulo com banco e conta
        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = subtitulo.add_run(f"{banco}  AG.{agencia}    C/C {conta}")
        run_sub.font.size = Pt(11)
        run_sub.font.bold = True
        run_sub.font.name = "Arial"

        # Nome do mês
        paragrafo_mes = doc.add_paragraph()
        paragrafo_mes.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_mes = paragrafo_mes.add_run(nome_mes)
        run_mes.font.size = Pt(12)
        run_mes.font.bold = True
        run_mes.font.name = "Arial"

        # Espaço
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Criar tabela
        num_linhas = len(lancs_mes) + 1  # +1 para cabeçalho
        tabela = doc.add_table(rows=num_linhas, cols=4)
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Definir larguras das colunas
        larguras = [Cm(2.5), Cm(2.0), Cm(10.0), Cm(3.0)]
        for row in tabela.rows:
            for idx, width in enumerate(larguras):
                row.cells[idx].width = width

        # Cabeçalho
        cabecalhos = ["DATA", "DEB/CRED", "DESCRIÇÃO", "VALOR"]
        for idx, texto in enumerate(cabecalhos):
            cell = tabela.rows[0].cells[idx]
            _formatar_celula(cell, texto, negrito=True, tamanho=9,
                           alinhamento=WD_ALIGN_PARAGRAPH.CENTER)

        # Preencher lançamentos
        for i, lanc in enumerate(lancs_mes):
            row_idx = i + 1
            row = tabela.rows[row_idx]

            data = lanc.get("data", "")
            tipo = lanc.get("tipo", "").upper()
            descricao = lanc.get("descricao", "")
            valor = lanc.get("valor", "")

            # Determinar se é crédito (deve ficar em negrito)
            is_credito = tipo in ("C", "CREDITO", "CRÉDITO")

            tipo_texto = "CRED" if is_credito else "DEB"

            _formatar_celula(row.cells[0], data, negrito=is_credito, tamanho=9,
                           alinhamento=WD_ALIGN_PARAGRAPH.CENTER)
            _formatar_celula(row.cells[1], tipo_texto, negrito=is_credito, tamanho=9,
                           alinhamento=WD_ALIGN_PARAGRAPH.CENTER)
            _formatar_celula(row.cells[2], descricao, negrito=is_credito, tamanho=9)
            _formatar_celula(row.cells[3], valor, negrito=is_credito, tamanho=9,
                           alinhamento=WD_ALIGN_PARAGRAPH.RIGHT)

        # Aplicar bordas na tabela
        tbl = tabela._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tblBorders>'
        )
        tblPr.append(borders)

    doc.save(output_path)
    return output_path
